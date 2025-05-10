#!/usr/bin/env python3
"""
Script to convert markdown manuscript to JAMA-compliant Word format.
Handles proper formatting, spacing, and page numbering according to JAMA guidelines.
"""

import os
import re
import sys
import logging
import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('jama_conversion.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger('jama_converter')

def add_page_numbers(doc):
    """Add page numbers to document in the upper right corner."""
    try:
        sections = doc.sections
        for section in sections:
            header = section.header
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            page_num = paragraph.add_run()
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            page_num._r.append(fldChar)
            page_num._r.append(instrText)
            page_num._r.append(fldChar2)
            
            # Set font properties for page numbers
            page_num.font.name = 'Times New Roman'
            page_num.font.size = Pt(12)
        
        logger.info("Page numbers added successfully")
    except Exception as e:
        logger.error(f"Error adding page numbers: {str(e)}")
        raise

def count_words(text):
    """Count words in text, excluding markdown formatting."""
    # Remove markdown formatting
    clean_text = clean_markdown(text)
    # Count words
    return len(re.findall(r'\S+', clean_text))

def clean_markdown(text):
    """Clean markdown formatting from text."""
    # Remove bold markers
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # Remove italic markers
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    # Remove reference links
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    # Remove superscript notation
    text = re.sub(r'\^(\d+)\^', r'\1', text)
    # Clean up any remaining markdown artifacts
    text = text.replace('\\', '')
    return text

def split_into_sections(content):
    """Split markdown content into labeled sections."""
    try:
        sections = {}
        
        # Extract title
        title_match = re.search(r'# (.*?)(?:\n|$)', content)
        if title_match:
            sections['title'] = title_match.group(1)
        
        # Extract sections using regex patterns
        section_patterns = {
            'title_page': r'## Title Page\s*\n(.*?)(?=\n##|\Z)',
            'abstract': r'## Abstract\s*\n(.*?)(?=\n##|\Z)', 
            'introduction': r'## (?:Background|Introduction)\s*\n(.*?)(?=\n##|\Z)',
            'methods': r'## Methods\s*\n(.*?)(?=\n##|\Z)',
            'results': r'## Results\s*\n(.*?)(?=\n##|\Z)',
            'discussion': r'## Discussion\s*\n(.*?)(?=\n##|\Z)',
            'conclusion': r'## Conclusions\s*\n(.*?)(?=\n##|\Z)',
            'abbreviations': r'## List of abbreviations\s*\n(.*?)(?=\n##|\Z)',
            'declarations': r'## Declarations\s*\n(.*?)(?=\n##|\Z)',
            'references': r'## References\s*\n(.*?)(?=\n##|\Z)',
        }
        
        for section_name, pattern in section_patterns.items():
            match = re.search(pattern, content, re.DOTALL)
            if match:
                sections[section_name] = match.group(1).strip()
        
        logger.info(f"Successfully extracted {len(sections)} sections from manuscript")
        return sections
    
    except Exception as e:
        logger.error(f"Error splitting markdown into sections: {str(e)}")
        raise

def process_title_page(doc, sections):
    """Process and format the title page according to JAMA requirements."""
    try:
        # Add title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_after = Pt(12)
        title_run = title_para.add_run(sections.get('title', 'Untitled Manuscript'))
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        title_run.font.name = 'Times New Roman'
        
        # Extract author information from title page
        if 'title_page' in sections:
            author_text = sections['title_page']
            
            # Process authors
            authors_match = re.search(r'\*\*Authors\*\*:\s*(.*?)(?=\n\n\*\*|\Z)', author_text, re.DOTALL)
            if authors_match:
                authors_para = doc.add_paragraph()
                authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                authors_para.paragraph_format.space_after = Pt(12)
                authors_run = authors_para.add_run(clean_markdown(authors_match.group(1)))
                authors_run.font.size = Pt(12)
                authors_run.font.name = 'Times New Roman'
            
            # Process affiliations
            affil_match = re.search(r'\*\*Affiliations\*\*:\s*(.*?)(?=\n\n\*\*|\Z)', author_text, re.DOTALL)
            if affil_match:
                affil_para = doc.add_paragraph()
                affil_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                affil_para.paragraph_format.space_after = Pt(12)
                affil_run = affil_para.add_run(clean_markdown(affil_match.group(1)))
                affil_run.font.size = Pt(12)
                affil_run.font.name = 'Times New Roman'
            
            # Process corresponding author
            corresp_match = re.search(r'\*\*Corresponding Author\*\*:\s*(.*?)(?=\n\n\*\*|\Z)', author_text, re.DOTALL)
            if corresp_match:
                corresp_para = doc.add_paragraph()
                corresp_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                corresp_para.paragraph_format.space_after = Pt(12)
                corresp_run = corresp_para.add_run(clean_markdown(corresp_match.group(1)))
                corresp_run.font.size = Pt(12)
                corresp_run.font.name = 'Times New Roman'
        
        # Add word count (JAMA requirement)
        word_count = 0
        for section in ['introduction', 'methods', 'results', 'discussion', 'conclusion']:
            if section in sections:
                word_count += count_words(sections[section])
        
        word_count_p = doc.add_paragraph()
        word_count_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        word_count_run = word_count_p.add_run(f"Word Count: {word_count}")
        word_count_run.font.size = Pt(12)
        word_count_run.font.name = 'Times New Roman'
        
        # Page break after title page
        doc.add_page_break()
        logger.info("Title page processed successfully")
        
    except Exception as e:
        logger.error(f"Error processing title page: {str(e)}")
        raise

def process_abstract_page(doc, sections):
    """Process and format the abstract page according to JAMA requirements."""
    try:
        if 'abstract' not in sections:
            logger.warning("Abstract section not found in manuscript")
            return
        
        # Add "Abstract" heading
        abstract_title = doc.add_paragraph()
        abstract_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        abstract_title_run = abstract_title.add_run("ABSTRACT")
        abstract_title_run.font.size = Pt(12)
        abstract_title_run.font.bold = True
        abstract_title_run.font.name = 'Times New Roman'
        
        # Abstract word count
        abstract_word_count = count_words(sections['abstract'])
        abstract_count_p = doc.add_paragraph()
        abstract_count_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        abstract_count_run = abstract_count_p.add_run(f"Word Count: {abstract_word_count}")
        abstract_count_run.font.size = Pt(12)
        abstract_count_run.font.name = 'Times New Roman'
        
        # Abstract content
        try:
            abstract_parts = re.findall(r'\*\*(.*?)\*\*:\s*(.*?)(?=\n\n\*\*|\Z)', sections['abstract'], re.DOTALL)
            for part, content in abstract_parts:
                para = doc.add_paragraph()
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                para.paragraph_format.space_after = Pt(0)
                part_run = para.add_run(f"{part}: ")
                part_run.font.bold = True
                part_run.font.size = Pt(12)
                part_run.font.name = 'Times New Roman'
                
                content_run = para.add_run(clean_markdown(content.strip()))
                content_run.font.size = Pt(12)
                content_run.font.name = 'Times New Roman'
        except Exception as e:
            logger.error(f"Error processing structured abstract: {str(e)}")
            # Add a simple paragraph if structured abstract parsing fails
            para = doc.add_paragraph()
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            para_run = para.add_run(clean_markdown(sections['abstract']))
            para_run.font.size = Pt(12)
            para_run.font.name = 'Times New Roman'
        
        # Page break after abstract
        doc.add_page_break()
        logger.info("Abstract page processed successfully")
        
    except Exception as e:
        logger.error(f"Error processing abstract page: {str(e)}")
        raise

def process_paragraphs(doc, sections):
    """Process paragraphs with proper formatting for each section."""
    try:
        # Main content sections
        for section_name in ['introduction', 'methods', 'results', 'discussion', 'conclusion']:
            if section_name in sections:
                heading = doc.add_paragraph()
                heading.paragraph_format.space_before = Pt(12)
                heading.paragraph_format.space_after = Pt(6)
                heading_run = heading.add_run(section_name.capitalize())
                heading_run.font.bold = True
                heading_run.font.size = Pt(12)
                heading_run.font.name = 'Times New Roman'
                
                # Split content into paragraphs and process each one
                paragraphs = sections[section_name].split('\n\n')
                for para_text in paragraphs:
                    para = doc.add_paragraph()
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.first_line_indent = Inches(0.5)
                    
                    # Process any inline references or formatting
                    cleaned_text = clean_markdown(para_text.strip())
                    para_run = para.add_run(cleaned_text)
                    para_run.font.size = Pt(12)
                    para_run.font.name = 'Times New Roman'

        # Process abbreviations section if present
        if 'abbreviations' in sections:
            process_abbreviations_section(doc, sections['abbreviations'])

        # Process declarations section if present
        if 'declarations' in sections:
            process_declarations_section(doc, sections['declarations'])

        # Process references section if present
        if 'references' in sections:
            process_references_section(doc, sections['references'])

        logger.info("Paragraphs processed successfully")
    except Exception as e:
        logger.error(f"Error processing paragraphs: {str(e)}")
        raise

def process_abbreviations_section(doc, content):
    """Process the abbreviations section with proper formatting."""
    try:
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
        heading_run = heading.add_run("List of Abbreviations")
        heading_run.font.bold = True
        heading_run.font.size = Pt(12)
        heading_run.font.name = 'Times New Roman'
        
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        para_run = para.add_run(clean_markdown(content))
        para_run.font.size = Pt(12)
        para_run.font.name = 'Times New Roman'
        
        logger.info("Abbreviations section processed successfully")
    except Exception as e:
        logger.error(f"Error processing abbreviations section: {str(e)}")
        raise

def process_declarations_section(doc, content):
    """Process the declarations section with proper formatting."""
    try:
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
        heading_run = heading.add_run("Declarations")
        heading_run.font.bold = True
        heading_run.font.size = Pt(12)
        heading_run.font.name = 'Times New Roman'
        
        # Split declarations by sections
        declaration_sections = re.findall(r'### (.*?)\s*\n\n(.*?)(?=\n\n###|\Z)', content, re.DOTALL)
        
        for title, content in declaration_sections:
            subheading = doc.add_paragraph()
            subheading.paragraph_format.space_before = Pt(10)
            subheading.paragraph_format.space_after = Pt(6)
            subheading.paragraph_format.left_indent = Inches(0.25)
            subheading_run = subheading.add_run(title)
            subheading_run.font.italic = True
            subheading_run.font.size = Pt(12)
            subheading_run.font.name = 'Times New Roman'
            
            para = doc.add_paragraph()
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.left_indent = Inches(0.5)
            para_run = para.add_run(clean_markdown(content.strip()))
            para_run.font.size = Pt(12)
            para_run.font.name = 'Times New Roman'
        
        logger.info("Declarations section processed successfully")
    except Exception as e:
        logger.error(f"Error processing declarations section: {str(e)}")
        raise

def process_references_section(doc, content):
    """Process the references section with proper formatting."""
    try:
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
        heading_run = heading.add_run("References")
        heading_run.font.bold = True
        heading_run.font.size = Pt(12)
        heading_run.font.name = 'Times New Roman'
        
        # Process each reference
        references = content.split('\n')
        for ref in references:
            if ref.strip():
                para = doc.add_paragraph()
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.left_indent = Inches(0.5)
                para.paragraph_format.hanging_indent = Inches(0.5)
                para_run = para.add_run(clean_markdown(ref.strip()))
                para_run.font.size = Pt(12)
                para_run.font.name = 'Times New Roman'
        
        logger.info("References section processed successfully")
    except Exception as e:
        logger.error(f"Error processing references section: {str(e)}")
        raise

def convert_cover_letter(input_file, output_dir):
    """Convert cover letter from markdown to JAMA-compliant Word format."""
    try:
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
        
        # Read input markdown file
        with open(input_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Create new Word document
        doc = Document()
        
        # Set up document properties
        section = doc.sections[0]
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        # Current date
        date_para = doc.add_paragraph()
        date_para.paragraph_format.space_after = Pt(24)
        date_run = date_para.add_run(datetime.datetime.now().strftime("%B %d, %Y"))
        date_run.font.name = 'Times New Roman'
        date_run.font.size = Pt(12)
        
        # Extract addressee information if present
        addressee_match = re.search(r'^\s*To:\s*(.*?)(?=\n\n|\Z)', content, re.MULTILINE | re.DOTALL)
        if addressee_match:
            addressee_para = doc.add_paragraph()
            addressee_para.paragraph_format.space_after = Pt(24)
            addressee_text = clean_markdown(addressee_match.group(1).strip())
            addressee_run = addressee_para.add_run(addressee_text)
            addressee_run.font.name = 'Times New Roman'
            addressee_run.font.size = Pt(12)
            
            # Remove addressee part from content
            content = content.replace(addressee_match.group(0), '', 1)
        
        # Extract subject/RE line if present
        subject_match = re.search(r'^\s*Re:[ \t]*(.+?)$|^\s*Subject:[ \t]*(.+?)$', content, re.MULTILINE)
        if subject_match:
            subject_text = subject_match.group(1) or subject_match.group(2)
            subject_para = doc.add_paragraph()
            subject_para.paragraph_format.space_after = Pt(24)
            subject_run = subject_para.add_run(f"Re: {clean_markdown(subject_text.strip())}")
            subject_run.font.bold = True
            subject_run.font.name = 'Times New Roman'
            subject_run.font.size = Pt(12)
            
            # Remove subject line from content
            content = content.replace(subject_match.group(0), '', 1)
            
        # Extract and clean content paragraphs
        paragraphs = re.split(r'\n\n+', content.strip())
        
        # Process each paragraph
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph()
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                para.paragraph_format.space_after = Pt(12)
                para.paragraph_format.first_line_indent = Inches(0.5)
                para_run = para.add_run(clean_markdown(para_text.strip()))
                para_run.font.name = 'Times New Roman'
                para_run.font.size = Pt(12)
        
        # Add signature block
        sig_para = doc.add_paragraph()
        sig_para.paragraph_format.space_before = Pt(24)
        sig_para.paragraph_format.space_after = Pt(18)
        sig_run = sig_para.add_run("Sincerely,")
        sig_run.font.name = 'Times New Roman'
        sig_run.font.size = Pt(12)
        
        # Extract author names from manuscript
        manuscript_path = os.path.join(os.path.dirname(input_file), "manuscript_equity_focus.md")
        try:
            with open(manuscript_path, 'r', encoding='utf-8') as f:
                manuscript_content = f.read()
                
            # Extract author names
            authors_match = re.search(r'\*\*Authors\*\*:\s*(.*?)(?=\n\n\*\*|\Z)', manuscript_content, re.DOTALL)
            if authors_match:
                authors_text = authors_match.group(1)
                # Extract individual authors
                author_list = re.findall(r'([^,^]+)(?:\^\d+\^)?(?:,|\Z)', authors_text)
                
                # Add each author signature
                for author in author_list:
                    author_name = author.strip()
                    author_sig = doc.add_paragraph()
                    author_sig.paragraph_format.space_after = Pt(0)
                    author_run = author_sig.add_run(author_name)
                    author_run.font.name = 'Times New Roman'
                    author_run.font.size = Pt(12)
        except Exception as e:
            # If can't extract authors, add a generic signature line
            author_sig = doc.add_paragraph()
            author_sig.paragraph_format.space_before = Pt(36)
            author_run = author_sig.add_run("_______________________")
            author_run.font.name = 'Times New Roman'
            author_run.font.size = Pt(12)
        
        # Save the document
        output_file = os.path.join(output_dir, 'cover_letter.docx')
        doc.save(output_file)
        
        logger.info(f"Cover letter successfully converted and saved to {output_file}")
        return output_file
    
    except Exception as e:
        logger.error(f"Error converting cover letter: {str(e)}")
        raise

def convert_to_jama_format(input_file, output_dir):
    """Main function to convert markdown to JAMA-compliant Word format."""
    try:
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
        
        # Read input markdown file
        with open(input_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Split content into sections
        sections = split_into_sections(content)
        
        # Create new Word document
        doc = Document()
        
        # Set up document properties
        section = doc.sections[0]
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        # Add page numbers
        add_page_numbers(doc)
        
        # Process title page
        process_title_page(doc, sections)
        
        # Process abstract
        process_abstract_page(doc, sections)
        
        # Process main content
        process_paragraphs(doc, sections)
        
        # Save the document
        output_file = os.path.join(output_dir, 'manuscript_jama.docx')
        doc.save(output_file)
        
        logger.info(f"Document successfully converted and saved to {output_file}")
        return output_file
    
    except Exception as e:
        logger.error(f"Error converting document: {str(e)}")
        raise

if __name__ == "__main__":
    try:
        # Set input and output paths
        script_dir = os.path.dirname(os.path.abspath(__file__))
        project_dir = os.path.dirname(script_dir)
        
        manuscript_input = os.path.join(project_dir, "manuscript", "manuscript_equity_focus.md")
        cover_letter_input = os.path.join(project_dir, "manuscript", "cover_letter_equity.md")
        output_dir = os.path.join(project_dir, "JAMA_submission", "manuscript")
        
        # Ensure output directory exists
        os.makedirs(output_dir, exist_ok=True)
        
        logger.info("Starting JAMA document conversion")
        
        # Convert the manuscript
        logger.info("Converting manuscript...")
        manuscript_output = convert_to_jama_format(manuscript_input, output_dir)
        print(f"Successfully created JAMA-formatted manuscript at: {manuscript_output}")
        
        # Convert the cover letter
        logger.info("Converting cover letter...")
        cover_letter_output = convert_cover_letter(cover_letter_input, output_dir)
        print(f"Successfully created JAMA-formatted cover letter at: {cover_letter_output}")
        
        logger.info("All documents converted successfully")
        
    except Exception as e:
        logger.error(f"Script execution failed: {str(e)}")
        print(f"Error: {str(e)}")
        sys.exit(1)
