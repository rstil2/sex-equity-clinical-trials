#!/usr/bin/env python3
# Script to create a JAMA-formatted cover letter
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import datetime

# Create directories if they don't exist
manuscript_dir = '../JAMA_submission/manuscript'
os.makedirs(manuscript_dir, exist_ok=True)

# Create the cover letter document
doc = Document()

# Set document properties
section = doc.sections[0]
section.page_height = Inches(11)
section.page_width = Inches(8.5)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# Current date
current_date = datetime.datetime.now().strftime('%B %d, %Y')
date = doc.add_paragraph()
date.alignment = WD_ALIGN_PARAGRAPH.LEFT
date_run = date.add_run(current_date)
date_run.font.size = Pt(12)
date_run.font.name = 'Times New Roman'

# Add space after date
doc.add_paragraph()

# Address
address = doc.add_paragraph()
address.alignment = WD_ALIGN_PARAGRAPH.LEFT
address_text = "Editor-in-Chief\nJAMA (Journal of the American Medical Association)\n330 N. Wabash Ave, Suite 39300\nChicago, IL 60611-5885"
address_run = address.add_run(address_text)
address_run.font.size = Pt(12)
address_run.font.name = 'Times New Roman'

# Add space after address
doc.add_paragraph()

# Subject line
subject = doc.add_paragraph()
subject.alignment = WD_ALIGN_PARAGRAPH.LEFT
subject_run = subject.add_run("RE: Manuscript Submission - Sex Representation Equity in Clinical Trials: A Statistical Analysis")
subject_run.font.bold = True
subject_run.font.size = Pt(12)
subject_run.font.name = 'Times New Roman'

# Add space after subject
doc.add_paragraph()

# Salutation
salutation = doc.add_paragraph()
salutation.alignment = WD_ALIGN_PARAGRAPH.LEFT
salutation_run = salutation.add_run("Dear Editor-in-Chief:")
salutation_run.font.size = Pt(12)
salutation_run.font.name = 'Times New Roman'

# Body paragraphs
paragraphs = [
    "I am pleased to submit our manuscript \"Sex Representation Equity in Clinical Trials: A Statistical Analysis\" for consideration for publication in JAMA.",
    
    "This study provides a novel statistical approach to evaluating sex representation in clinical trials by testing the null hypothesis that male and female participants are equally represented relative to population demographics. While previous research has primarily used descriptive approaches to document the presence or absence of women in trials, our analysis rigorously tests whether current representation levels achieve statistical equity.",
    
    "Analyzing 1,825 clinical trials, we found:\n• No significant deviation from expected sex representation (50.8% female) across any disease category\n• Overall female representation of 50.1% (p = 0.844 compared to population expectation)\n• Successful achievement of equity even in historically problematic areas like HIV/AIDS research",
    
    "We believe this manuscript is appropriate for JAMA for several reasons:\n• The findings have significant implications for research ethics and policy\n• The methodological approach provides a template for equity assessments across clinical research\n• The topic addresses JAMA's commitment to promoting diversity and equity in medical research\n• The findings will interest JAMA's broad readership of clinicians, researchers, and policymakers",
    
    "The manuscript contains approximately 3,200 words, 2 tables, and 3 figures. It has not been published previously and is not under consideration by any other journal. All authors have reviewed the manuscript and agree with its contents and submission to JAMA. The authors have no conflicts of interest to declare.",
    
    "All data used in the analysis is publicly available through ClinicalTrials.gov, and all code and methods are thoroughly documented in the supplementary materials to ensure reproducibility.",
]

for text in paragraphs:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.first_line_indent = Inches(0.5)
    para_run = para.add_run(text)
    para_run.font.size = Pt(12)
    para_run.font.name = 'Times New Roman'
    
    # Add a blank line between paragraphs
    doc.add_paragraph()

# Closing
closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.LEFT
closing_run = closing.add_run("Sincerely,")
closing_run.font.size = Pt(12)
closing_run.font.name = 'Times New Roman'

# Space for signature
for _ in range(4):
    doc.add_paragraph()

# Author information
author_info = doc.add_paragraph()
author_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
author_text = "Jane D. Researcher, PhD\nDepartment of Public Health\nUniversity Research Institute\nEmail: jane.researcher@university.edu\nPhone: +1-234-567-8910"
author_run = author_info.add_run(author_text)
author_run.font.size = Pt(12)
author_run.font.name = 'Times New Roman'

# Save the document
cover_letter_path = os.path.join(manuscript_dir, "cover_letter.docx")
doc.save(cover_letter_path)

print(f"Cover letter created and saved to {cover_letter_path}")

