#!/usr/bin/env python3
# Script to create JAMA submission form templates
# Creates individual files for each author
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import datetime
import logging

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Create directories if they don't exist
forms_dir = '../JAMA_submission/forms'
os.makedirs(forms_dir, exist_ok=True)

# Author information from the manuscript
authors = [
    {
        'name': 'Jane D. Researcher',
        'affiliation': 'Department of Public Health, University Research Institute',
        'email': 'jane.researcher@university.edu',
        'orcid': '0000-0001-2345-6789'
    },
    {
        'name': 'Alex L. Scientist',
        'affiliation': 'Department of Public Health, University Research Institute',
        'email': 'alex.scientist@university.edu',
        'orcid': '0000-0002-3456-7890'
    },
    {
        'name': 'Morgan T. Analyst',
        'affiliation': 'Clinical Trials Research Center, University Medical School',
        'email': 'morgan.analyst@medical.edu',
        'orcid': '0000-0003-4567-8901'
    }
]

# Manuscript details
manuscript_title = "Sex Representation Equity in Clinical Trials: A Statistical Analysis"
manuscript_id = "JAMA-" + datetime.datetime.now().strftime('%Y-%m-%d')

# Function to create the author contribution form for a single author
def create_author_contribution_form(author):
    doc = Document()
    
    # Extract last name for file naming
    last_name = author['name'].split()[-1]
    
    # Set document properties
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("JAMA AUTHOR CONTRIBUTION FORM")
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.name = 'Times New Roman'
    
    # Add manuscript information
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run(f"Manuscript Title: {manuscript_title}\nManuscript ID: {manuscript_id}")
    info_run.font.size = Pt(12)
    info_run.font.name = 'Times New Roman'
    
    # Add author information
    author_info = doc.add_paragraph()
    author_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_info_run = author_info.add_run(f"Author: {author['name']}\nAffiliation: {author['affiliation']}\nEmail: {author['email']}")
    author_info_run.font.size = Pt(12)
    author_info_run.font.name = 'Times New Roman'
    
    # Add instructions
    instructions = doc.add_paragraph()
    instructions.paragraph_format.space_before = Pt(12)
    instructions.paragraph_format.space_after = Pt(12)
    instructions_run = instructions.add_run("The International Committee of Medical Journal Editors (ICMJE) recommends that authorship be based on the following 4 criteria:")
    instructions_run.font.size = Pt(12)
    instructions_run.font.name = 'Times New Roman'
    
    # Add criteria
    criteria = [
        "Substantial contributions to the conception or design of the work; or the acquisition, analysis, or interpretation of data for the work; AND",
        "Drafting the work or revising it critically for important intellectual content; AND",
        "Final approval of the version to be published; AND",
        "Agreement to be accountable for all aspects of the work in ensuring that questions related to the accuracy or integrity of any part of the work are appropriately investigated and resolved."
    ]
    
    for criterion in criteria:
        item = doc.add_paragraph()
        item.paragraph_format.left_indent = Inches(0.5)
        item.paragraph_format.space_after = Pt(6)
        item_run = item.add_run(criterion)
        item_run.font.size = Pt(12)
        item_run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    note.paragraph_format.space_after = Pt(12)
    note_run = note.add_run("Please check the appropriate boxes below and sign at the bottom.")
    note_run.font.size = Pt(12)
    note_run.font.italic = True
    note_run.font.name = 'Times New Roman'
    
    # Create contribution checklist
    contribution_categories = [
        "Conception or design",
        "Acquisition of data",
        "Analysis or interpretation of data",
        "Drafting the manuscript",
        "Critical revision of the manuscript",
        "Statistical analysis",
        "Obtaining funding",
        "Administrative or technical support",
        "Supervision"
    ]
    
    for category in contribution_categories:
        contribution = doc.add_paragraph()
        contribution.paragraph_format.left_indent = Inches(0.5)
        contribution.paragraph_format.space_after = Pt(6)
        contribution_run = contribution.add_run(f"□ {category}")
        contribution_run.font.size = Pt(12)
        contribution_run.font.name = 'Times New Roman'
    
    # Signature line
    signature = doc.add_paragraph()
    signature.paragraph_format.space_before = Pt(24)
    signature.paragraph_format.space_after = Pt(12)
    signature_run = signature.add_run("Signature: __________________________________ Date: _______________")
    signature_run.font.size = Pt(12)
    signature_run.font.name = 'Times New Roman'
    
    # Save the document with author-specific filename
    file_path = os.path.join(forms_dir, f'author_contribution_{last_name}.docx')
    doc.save(file_path)
    logger.info(f"Author Contribution Form created for {last_name}")
    print(f"Author Contribution Form created for {last_name}")
    return file_path

# Function to create the ICMJE disclosure form for a single author
def create_icmje_disclosure_form(author):
    """Create ICMJE disclosure form for a single author."""
    doc = Document()
    
    # Extract last name for file naming
    last_name = author['name'].split()[-1]
    
    # Set document properties
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ICMJE DISCLOSURE FORM")
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.name = 'Times New Roman'
    
    # Add manuscript information
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run(f"Manuscript Title: {manuscript_title}\nManuscript ID: {manuscript_id}")
    info_run.font.size = Pt(12)
    info_run.font.name = 'Times New Roman'
    
    # Add author information
    author_info = doc.add_paragraph()
    author_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_info_run = author_info.add_run(f"Author: {author['name']}\nAffiliation: {author['affiliation']}\nEmail: {author['email']}")
    author_info_run.font.size = Pt(12)
    author_info_run.font.name = 'Times New Roman'
    
    # Add disclosure sections
    disclosure_sections = [
        "1. Financial relationships with industry",
        "2. Academic or institutional affiliations",
        "3. Research support or funding",
        "4. Intellectual property rights",
        "5. Personal relationships",
        "6. Other potential conflicts of interest"
    ]
    
    for section in disclosure_sections:
        # Section title
        section_title = doc.add_paragraph()
        section_title.paragraph_format.space_before = Pt(12)
        section_title.paragraph_format.space_after = Pt(6)
        section_title_run = section_title.add_run(section)
        section_title_run.font.size = Pt(12)
        section_title_run.font.name = 'Times New Roman'
        
        # Response options
        response = doc.add_paragraph()
        response.paragraph_format.left_indent = Inches(0.5)
        response.paragraph_format.space_after = Pt(12)
        response_run = response.add_run("□ No\n□ Yes (explain below)")
        response_run.font.size = Pt(12)
        response_run.font.name = 'Times New Roman'
        
        # Explanation space
        explanation = doc.add_paragraph()
        explanation.paragraph_format.left_indent = Inches(0.5)
        explanation.paragraph_format.space_after = Pt(12)
        explanation_run = explanation.add_run("If yes, please explain: _____________________________________________")
        explanation_run.font.size = Pt(12)
        explanation_run.font.name = 'Times New Roman'
    
    # Signature line
    signature = doc.add_paragraph()
    signature.paragraph_format.space_before = Pt(24)
    signature.paragraph_format.space_after = Pt(12)
    signature_run = signature.add_run("Signature: __________________________________ Date: _______________")
    signature_run.font.size = Pt(12)
    signature_run.font.name = 'Times New Roman'
    
    # Save the document with author-specific filename
    file_path = os.path.join(forms_dir, f'icmje_disclosure_{last_name}.docx')
    doc.save(file_path)
    logger.info(f"ICMJE Disclosure Form created for {last_name}")
    print(f"ICMJE Disclosure Form created for {last_name}")
    return file_path

# Function to create the copyright transfer agreement
def create_copyright_agreement():
    doc = Document()
    
    # Set document properties
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("COPYRIGHT TRANSFER AGREEMENT")
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.name = 'Times New Roman'
    
    # Add manuscript information
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run(f"Manuscript Title: {manuscript_title}\nManuscript ID: {manuscript_id}")
    info_run.font.size = Pt(12)
    info_run.font.name = 'Times New Roman'
    
    # Add agreement text
    agreement_text = (
        "The undersigned authors hereby transfer, assign, or otherwise convey all copyright "
        "ownership, including any and all rights incidental thereto, exclusively to the "
        "American Medical Association (AMA), in the event that such work is published in JAMA. "
        "This agreement is binding on the authors, their heirs, and the AMA.\n\n"
        "The authors represent and warrant that:\n"
        "1. The manuscript is original, has not been previously published, and is not under consideration "
        "for publication elsewhere.\n"
        "2. They are the sole authors and owners of the manuscript and have full authority to enter into "
        "this agreement.\n"
        "3. The manuscript does not infringe upon any copyright, proprietary right, or any other right "
        "of any third party.\n"
        "4. The manuscript contains no material that is defamatory, violates any right of privacy, or is "
        "otherwise contrary to law.\n"
        "5. They have made a significant scientific contribution to the study and approved the final version.\n"
        "6. If the manuscript was prepared jointly with other authors, they have informed the co-author(s) "
        "of the terms of this agreement and are signing on their behalf.\n\n"
        "The authors understand that if the manuscript is accepted for publication, they will be required to "
        "pay the publication fees as determined by the journal."
    )
    
    agreement = doc.add_paragraph()
    agreement.paragraph_format.space_before = Pt(12)
    agreement.paragraph_format.space_after = Pt(12)
    agreement.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    agreement_run = agreement.add_run(agreement_text)
    agreement_run.font.size = Pt(12)
    agreement_run.font.name = 'Times New Roman'
    
    # Add signature sections for each author
    for author in authors:
        # Author name and signature line
        signature = doc.add_paragraph()
        signature.paragraph_format.space_before = Pt(12)
        signature.paragraph_format.space_after = Pt(6)
        signature_text = f"Author Name: {author['name']}\nAffiliation: {author['affiliation']}\nEmail: {author['email']}"
        signature_run = signature.add_run(signature_text)
        signature_run.font.size = Pt(12)
        signature_run.font.name = 'Times New Roman'
        
        # Signature line
        sig_line = doc.add_paragraph()
        sig_line.paragraph_format.space_before = Pt(24)
        sig_line.paragraph_format.space_after = Pt(12)
        sig_line_run = sig_line.add_run("Signature: __________________________________ Date: _______________")
        sig_line_run.font.size = Pt(12)
        sig_line_run.font.name = 'Times New Roman'
        
        # Page break between authors (except for the last author)
        if author != authors[-1]:
            doc.add_page_break()
    
    # Save the document
    file_path = os.path.join(forms_dir, 'copyright_transfer_agreement.docx')
    doc.save(file_path)
    logger.info(f"Created copyright transfer agreement")
    print(f"Copyright Transfer Agreement created and saved to {file_path}")
    return file_path

# Execute form creation functions for each author
def main():
    try:
        logger.info("Starting JAMA forms creation process")
        
        # Clean up existing files
        logger.info("Cleaning up existing files")
        for file in os.listdir(forms_dir):
            if file.endswith('.docx'):
                file_path = os.path.join(forms_dir, file)
                os.remove(file_path)
                logger.info(f"Removed existing file: {file}")
        
        created_files = []
        
        # Create individual forms for each author
        for author in authors:
            try:
                last_name = author['name'].split()[-1]
                logger.info(f"Creating forms for {last_name}")
                
                # Create author contribution form
                contribution_file = create_author_contribution_form(author)
                created_files.append(os.path.basename(contribution_file))
                
                # Create ICMJE disclosure form
                disclosure_file = create_icmje_disclosure_form(author)
                created_files.append(os.path.basename(disclosure_file))
                
                logger.info(f"Successfully created forms for {last_name}")
            except Exception as e:
                logger.error(f"Error creating forms for {author['name']}: {str(e)}")
                raise
        
        # Create single copyright agreement for all authors
        copyright_file = create_copyright_agreement()
        created_files.append(os.path.basename(copyright_file))
        
        # Verify all expected files were created
        expected_files = set([
            f"author_contribution_{author['name'].split()[-1]}.docx" for author in authors
        ] + [
            f"icmje_disclosure_{author['name'].split()[-1]}.docx" for author in authors
        ] + ['copyright_transfer_agreement.docx'])
        
        actual_files = set(os.listdir(forms_dir))
        
        if expected_files != actual_files:
            missing = expected_files - actual_files
            extra = actual_files - expected_files
            error_msg = ""
            if missing:
                error_msg += f"Missing files: {missing}\n"
            if extra:
                error_msg += f"Unexpected files: {extra}\n"
            logger.error(f"File verification failed: {error_msg}")
            raise Exception(f"File verification failed: {error_msg}")
        
        logger.info("Successfully completed JAMA forms creation")
        logger.info(f"Created {len(expected_files)} files in {forms_dir}")
    except Exception as e:
        logger.error(f"Error in JAMA forms creation process: {str(e)}")
        raise

if __name__ == "__main__":
    main()

