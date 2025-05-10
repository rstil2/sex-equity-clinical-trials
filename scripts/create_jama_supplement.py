#!/usr/bin/env python3
# Script to create JAMA-formatted supplementary materials
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

# Create directories if they don't exist
manuscript_dir = '../JAMA_submission/manuscript'
os.makedirs(manuscript_dir, exist_ok=True)

# Function to read the markdown file
def read_markdown_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

# Function to convert Python code blocks to formatted text
def format_code_block(code_text):
    # Remove the triple backticks
    code_text = re.sub(r'^```python\s*$', '', code_text, flags=re.MULTILINE)
    code_text = re.sub(r'^```\s*$', '', code_text, flags=re.MULTILINE)
    
    # Clean up any extra whitespace
    return code_text.strip()

# Function to extract sections from the markdown content
def extract_sections(content):
    sections = {}
    
    # Extract main title
    title_match = re.search(r'# (.*)', content)
    if title_match:
        sections['title'] = title_match.group(1)
    
    # Extract detailed data processing section
    data_processing_match = re.search(r'## Detailed Data Processing and Analysis Methods\s*\n\n(.*?)(?=\n\n## Software)', content, re.DOTALL)
    if data_processing_match:
        sections['data_processing'] = data_processing_match.group(1)
    
    # Extract software versions section
    software_match = re.search(r'## Software and Package Versions\s*\n\n(.*?)$', content, re.DOTALL)
    if software_match:
        sections['software'] = software_match.group(1)
    
    # Extract code blocks
    code_blocks = re.findall(r'```python\s*(.*?)```', content, re.DOTALL)
    if code_blocks:
        sections['code_blocks'] = code_blocks
        
    # Extract section titles and content
    section_matches = re.findall(r'### (.*?)\s*\n\n(.*?)(?=\n\n###|\n\n##|\Z)', content, re.DOTALL)
    sections['subsections'] = section_matches
    
    return sections

# Create JAMA-formatted supplementary materials document
def create_jama_supplement(source_path, output_path):
    # Read the markdown content
    content = read_markdown_file(source_path)
    
    # Extract sections
    sections = extract_sections(content)
    
    # Create a new document
    doc = Document()
    
    # Set document properties
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Supplementary Materials")
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.name = 'Times New Roman'
    
    # Add subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Sex Representation Equity in Clinical Trials: A Statistical Analysis")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True
    subtitle_run.font.name = 'Times New Roman'
    
    # Add a page break after the title page
    doc.add_page_break()
    
    # Add the "Detailed Methods" section heading
    detailed_methods_heading = doc.add_paragraph()
    detailed_methods_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    detailed_methods_heading_run = detailed_methods_heading.add_run("DETAILED METHODS")
    detailed_methods_heading_run.font.size = Pt(12)
    detailed_methods_heading_run.font.bold = True
    detailed_methods_heading_run.font.name = 'Times New Roman'
    
    # Process subsections
    if 'subsections' in sections:
        for title, content in sections['subsections']:
            # Add subsection heading
            heading = doc.add_paragraph()
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
            heading_run = heading.add_run(title)
            heading_run.font.bold = True
            heading_run.font.size = Pt(12)
            heading_run.font.name = 'Times New Roman'
            
            # Check if this section contains code
            if "```python" in content:
                # Split by code blocks
                parts = re.split(r'(```python.*?```)', content, flags=re.DOTALL)
                for part in parts:
                    if part.strip().startswith("```python"):
                        # This is a code block
                        code_text = format_code_block(part)
                        
                        # Add the code with a different style
                        code_para = doc.add_paragraph()
                        code_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        code_para.paragraph_format.space_after = Pt(6)
                        code_para.paragraph_format.left_indent = Inches(0.5)
                        code_run = code_para.add_run(code_text)
                        code_run.font.size = Pt(10)
                        code_run.font.name = 'Courier New'
                    else:
                        # This is regular text
                        if part.strip():
                            text_para = doc.add_paragraph()
                            text_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                            text_para.paragraph_format.space_after = Pt(0)
                            text_run = text_para.add_run(part.strip())
                            text_run.font.size = Pt(12)
                            text_run.font.name = 'Times New Roman'
            else:
                # Regular text section
                text_para = doc.add_paragraph()
                text_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                text_para.paragraph_format.space_after = Pt(0)
                text_run = text_para.add_run(content.strip())
                text_run.font.size = Pt(12)
                text_run.font.name = 'Times New Roman'
    
    # Add software versions section
    if 'software' in sections:
        software_heading = doc.add_paragraph()
        software_heading.paragraph_format.space_before = Pt(12)
        software_heading.paragraph_format.space_after = Pt(6)
        software_heading_run = software_heading.add_run("Software and Package Versions")
        software_heading_run.font.bold = True
        software_heading_run.font.size = Pt(12)
        software_heading_run.font.name = 'Times New Roman'
        
        # Create a formatted list
        lines = sections['software'].strip().split('\n')
        for line in lines:
            if line.strip():
                item_para = doc.add_paragraph()
                item_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                item_para.paragraph_format.left_indent = Inches(0.25)
                item_para.paragraph_format.space_after = Pt(0)
                item_run = item_para.add_run(line.strip())
                item_run.font.size = Pt(12)
                item_run.font.name = 'Times New Roman'
    
    # Save the document
    doc.save(output_path)
    print(f"Supplementary Materials document created and saved to {output_path}")

# Main execution
source_path = '../manuscript/supplementary_methods.md'
output_path = os.path.join(manuscript_dir, 'supplementary_materials.docx')
create_jama_supplement(source_path, output_path)

