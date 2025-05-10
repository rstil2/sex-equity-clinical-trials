#!/usr/bin/env python3
# Script to create JAMA-formatted tables
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

# Create directories if they don't exist
tables_dir = '../JAMA_submission/tables'
os.makedirs(tables_dir, exist_ok=True)

# Function to read a markdown file
def read_markdown_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

# Function to parse markdown table into rows and columns
def parse_markdown_table(markdown_content):
    # Extract title
    title_match = re.search(r'\*\*Table \d+: (.*?)\*\*', markdown_content)
    title = title_match.group(1) if title_match else "Table"
    
    # Extract footnote
    footnote_match = re.search(r'Note: (.*?)$', markdown_content, re.MULTILINE)
    footnote = footnote_match.group(1) if footnote_match else ""
    
    # Extract table content
    lines = markdown_content.strip().split('\n')
    rows = []
    
    # Find lines that contain table rows (those with pipes)
    for line in lines:
        if '|' in line and not line.startswith('Note:'):
            # Skip separator rows (those with dashes)
            if re.match(r'^\s*\|[-\s|]+\|\s*$', line):
                continue
                
            # Extract cells from the row
            cells = []
            for cell in line.split('|'):
                if cell.strip():  # Skip empty cells at the beginning/end
                    # Clean up cell content - remove markdown formatting
                    clean_cell = re.sub(r'\*\*(.*?)\*\*', r'\1', cell.strip())
                    cells.append(clean_cell)
            
            if cells:  # Add only non-empty rows
                rows.append(cells)
    
    return {
        'title': title,
        'footnote': footnote,
        'rows': rows
    }

# Function to create a JAMA-formatted table document
def create_jama_table_document(table_data, table_number):
    doc = Document()
    
    # Set document properties
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"Table {table_number}. {table_data['title']}")
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    title_run.font.name = 'Times New Roman'
    
    # Create table
    num_rows = len(table_data['rows'])
    if num_rows > 0:
        num_cols = len(table_data['rows'][0])
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Fill in table cells
        for i, row in enumerate(table_data['rows']):
            for j, cell_text in enumerate(row):
                if j < num_cols:  # Ensure we don't exceed column count
                    cell = table.cell(i, j)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    
                    # Set cell properties
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # For header row, make text bold
                    if i == 0:
                        run = paragraph.add_run(cell_text)
                        run.font.bold = True
                        run.font.size = Pt(12)
                        run.font.name = 'Times New Roman'
                        
                        # Add light gray background to header cells
                        # Note: python-docx doesn't directly support cell background colors
                        # This would typically be set through table styles in a real implementation
                    else:
                        # For numeric cells, right-align
                        if re.match(r'^-?\d+(\.\d+)?%?$', cell_text.strip()) or cell_text.strip() in ['p', 'P']:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        # For the first column (disease categories), left-align
                        elif j == 0:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        run = paragraph.add_run(cell_text)
                        run.font.size = Pt(12)
                        run.font.name = 'Times New Roman'
    
    # Add footnote
    if table_data['footnote']:
        footnote = doc.add_paragraph()
        footnote.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footnote_run = footnote.add_run(f"Note: {table_data['footnote']}")
        footnote_run.font.size = Pt(10)
        footnote_run.font.italic = True
        footnote_run.font.name = 'Times New Roman'
    
    return doc

# Process Table 1
table1_content = read_markdown_file('../tables/table1.md')
table1_data = parse_markdown_table(table1_content)
table1_doc = create_jama_table_document(table1_data, 1)
table1_path = os.path.join(tables_dir, 'table1.docx')
table1_doc.save(table1_path)
print(f"Table 1 created and saved to {table1_path}")

# Process Table 2
table2_content = read_markdown_file('../tables/table2.md')
table2_data = parse_markdown_table(table2_content)
table2_doc = create_jama_table_document(table2_data, 2)
table2_path = os.path.join(tables_dir, 'table2.docx')
table2_doc.save(table2_path)
print(f"Table 2 created and saved to {table2_path}")

